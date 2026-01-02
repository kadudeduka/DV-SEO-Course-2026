#!/usr/bin/env python3
"""
Lab Submission Markdown to DOCX Template Parser

This script parses lab submission markdown files and generates structured DOCX templates
for learner submissions. The templates include read-only sections, editable answer areas,
and trainer feedback sections.

Usage:
    python parse_lab_submission.py <submission_markdown_file> [output_dir]

Example:
    python parse_lab_submission.py data/courses/seo-master-2026/content/labs/Day_01_Lab_01_Submission_Format.md labs/templates/
"""

import re
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown


class LabSubmissionParser:
    """Parser for lab submission markdown files"""
    
    def __init__(self, markdown_file):
        self.markdown_file = Path(markdown_file)
        self.content = self._read_markdown()
        self.metadata = {}
        self.sections = []
        
    def _read_markdown(self):
        """Read markdown file content"""
        if not self.markdown_file.exists():
            raise FileNotFoundError(f"Markdown file not found: {self.markdown_file}")
        
        with open(self.markdown_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    def parse(self):
        """Parse markdown content into structured sections"""
        # Extract metadata
        self._extract_metadata()
        
        # Extract sections
        self._extract_sections()
        
        return {
            'metadata': self.metadata,
            'sections': self.sections
        }
    
    def _extract_metadata(self):
        """Extract metadata from markdown header"""
        # Extract Day and Lab Title from first line
        first_line_match = re.search(r'#\s*Lab Submission\s*—\s*Day\s*(\d+),\s*Lab\s*(\d+)', self.content, re.IGNORECASE)
        if first_line_match:
            self.metadata['day'] = int(first_line_match.group(1))
            self.metadata['lab_number'] = int(first_line_match.group(2))
        
        # Extract Lab Title
        lab_title_match = re.search(r'\*\*Lab Title:\*\*\s*(.+?)(?:\n|\*\*)', self.content)
        if lab_title_match:
            self.metadata['lab_title'] = lab_title_match.group(1).strip()
        
        # Extract Day Title
        day_title_match = re.search(r'\*\*Day:\*\*\s*(.+?)(?:\n|\*\*)', self.content)
        if day_title_match:
            self.metadata['day_title'] = day_title_match.group(1).strip()
        
        # Default course name
        self.metadata['course_name'] = 'SEO Master Course 2026'
    
    def _extract_sections(self):
        """Extract all sections from markdown"""
        # Split by horizontal rules (---)
        parts = re.split(r'\n---\n', self.content)
        
        current_section = None
        
        for part in parts:
            # Check for section headers (##)
            section_match = re.search(r'^##\s+(.+)$', part, re.MULTILINE)
            if section_match:
                if current_section:
                    self.sections.append(current_section)
                
                current_section = {
                    'title': section_match.group(1).strip(),
                    'type': self._classify_section(section_match.group(1)),
                    'content': part
                }
            elif current_section:
                current_section['content'] += '\n---\n' + part
        
        if current_section:
            self.sections.append(current_section)
    
    def _classify_section(self, title):
        """Classify section type based on title"""
        title_lower = title.lower()
        
        if 'trainer' in title_lower or 'feedback' in title_lower or 'score' in title_lower:
            return 'trainer'
        elif 'self-assessment' in title_lower or 'reflection' in title_lower:
            return 'self_assessment'
        elif 'instruction' in title_lower or 'context' in title_lower:
            return 'instruction'
        elif 'checklist' in title_lower or 'declaration' in title_lower:
            return 'meta'
        else:
            return 'question'


class DOCXGenerator:
    """Generate DOCX templates from parsed lab submission data"""
    
    def __init__(self, parsed_data):
        self.metadata = parsed_data['metadata']
        self.sections = parsed_data['sections']
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Setup document styles and properties"""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def generate(self):
        """Generate complete DOCX document"""
        self._add_cover_section()
        self._add_instructions_section()
        self._add_question_sections()
        self._add_self_assessment_sections()
        self._add_trainer_feedback_section()
        self._add_submission_declaration()
        
        return self.doc
    
    def _add_cover_section(self):
        """Add cover section with metadata"""
        # Title
        title = self.doc.add_heading('Lab Submission Template', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Course Name
        self.doc.add_paragraph().add_run('Course: ').bold = True
        self.doc.add_paragraph(self.metadata.get('course_name', 'SEO Master Course 2026'))
        
        # Day
        self.doc.add_paragraph().add_run('Day: ').bold = True
        self.doc.add_paragraph(self.metadata.get('day_title', f"Day {self.metadata.get('day', 'N/A')}"))
        
        # Lab Title
        self.doc.add_paragraph().add_run('Lab Title: ').bold = True
        self.doc.add_paragraph(self.metadata.get('lab_title', 'N/A'))
        
        # Learner Information (editable fields)
        self.doc.add_paragraph()
        learner_name_p = self.doc.add_paragraph()
        learner_name_p.add_run('Learner Name: ').bold = True
        learner_name_run = learner_name_p.add_run('_________________________')
        self._make_editable(learner_name_run)
        
        email_p = self.doc.add_paragraph()
        email_p.add_run('Email: ').bold = True
        email_run = email_p.add_run('_________________________')
        self._make_editable(email_run)
        
        # Submission Attempt
        attempt_p = self.doc.add_paragraph()
        attempt_p.add_run('Submission Attempt: ').bold = True
        attempt_p.add_run('1 (Auto-incremented on resubmission)')
        
        # Date
        date_p = self.doc.add_paragraph()
        date_p.add_run('Date: ').bold = True
        date_run = date_p.add_run('_________________________')
        self._make_editable(date_run)
        
        self.doc.add_page_break()
    
    def _add_instructions_section(self):
        """Add instructions section (read-only)"""
        instruction_sections = [s for s in self.sections if s['type'] == 'instruction']
        
        if instruction_sections:
            self.doc.add_heading('Instructions', 1)
            
            for section in instruction_sections:
                # Parse markdown content and convert to DOCX
                self._add_markdown_content(section['content'], read_only=True)
            
            self.doc.add_page_break()
    
    def _add_question_sections(self):
        """Add question sections with editable answer areas"""
        question_sections = [s for s in self.sections if s['type'] == 'question']
        
        for section in question_sections:
            # Add section title
            self.doc.add_heading(section['title'], 1)
            
            # Parse and add content
            self._add_markdown_content(section['content'], read_only=False)
            
            self.doc.add_paragraph()  # Spacing
    
    def _add_self_assessment_sections(self):
        """Add self-assessment sections"""
        assessment_sections = [s for s in self.sections if s['type'] == 'self_assessment']
        
        if assessment_sections:
            self.doc.add_page_break()
            self.doc.add_heading('Self-Assessment', 1)
            
            for section in assessment_sections:
                self._add_markdown_content(section['content'], read_only=False)
    
    def _add_trainer_feedback_section(self):
        """Add trainer feedback section (locked)"""
        self.doc.add_page_break()
        self.doc.add_heading('Trainer Feedback (Read-Only)', 1)
        
        # Feedback field
        feedback_p = self.doc.add_paragraph()
        feedback_p.add_run('Feedback: ').bold = True
        feedback_run = feedback_p.add_run('\n[Trainer feedback will appear here]')
        self._make_read_only(feedback_run)
        
        # Status field
        status_p = self.doc.add_paragraph()
        status_p.add_run('Status: ').bold = True
        status_run = status_p.add_run('[Pending / Approved / Needs Revision]')
        self._make_read_only(status_run)
        
        # Score field
        score_p = self.doc.add_paragraph()
        score_p.add_run('Score: ').bold = True
        score_run = score_p.add_run('[Score will appear here]')
        self._make_read_only(score_run)
    
    def _add_submission_declaration(self):
        """Add submission declaration section"""
        meta_sections = [s for s in self.sections if s['type'] == 'meta']
        
        if meta_sections:
            self.doc.add_page_break()
            
            for section in meta_sections:
                if 'declaration' in section['title'].lower():
                    self._add_markdown_content(section['content'], read_only=False)
    
    def _add_markdown_content(self, markdown_text, read_only=False):
        """Convert markdown content to DOCX paragraphs"""
        # Remove horizontal rules
        markdown_text = re.sub(r'^---\s*$', '', markdown_text, flags=re.MULTILINE)
        
        # Split into lines
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Handle headers
            if line.startswith('###'):
                text = line[3:].strip()
                p = self.doc.add_heading(text, 3)
                if read_only:
                    self._make_read_only_paragraph(p)
            elif line.startswith('##'):
                text = line[2:].strip()
                p = self.doc.add_heading(text, 2)
                if read_only:
                    self._make_read_only_paragraph(p)
            elif line.startswith('#'):
                text = line[1:].strip()
                p = self.doc.add_heading(text, 1)
                if read_only:
                    self._make_read_only_paragraph(p)
            # Handle bold text
            elif line.startswith('**') and line.endswith('**'):
                text = line[2:-2].strip()
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                if read_only:
                    self._make_read_only_paragraph(p)
            # Handle checkboxes
            elif line.startswith('- [ ]'):
                text = line[5:].strip()
                p = self.doc.add_paragraph(text, style='List Bullet')
                if not read_only:
                    self._make_editable_paragraph(p)
            # Handle regular paragraphs
            else:
                # Clean up markdown formatting
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)  # Remove bold
                text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Remove links
                
                if text:
                    p = self.doc.add_paragraph(text)
                    if read_only:
                        self._make_read_only_paragraph(p)
                    else:
                        # Check if this looks like an answer area
                        if re.search(r'\[.*\]|___+|\(.*\)', text):
                            self._make_editable_paragraph(p)
            
            i += 1
    
    def _make_editable(self, run):
        """Make a run editable (highlighted background)"""
        # Add light yellow background
        run.font.highlight_color = 7  # Yellow highlight
    
    def _make_editable_paragraph(self, paragraph):
        """Make entire paragraph editable with highlight"""
        for run in paragraph.runs:
            run.font.highlight_color = 7  # Yellow highlight
    
    def _make_read_only(self, run):
        """Make a run read-only (grayed out)"""
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True
    
    def _make_read_only_paragraph(self, paragraph):
        """Make entire paragraph read-only"""
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)


def generate_docx_template(markdown_file, output_dir=None, course_name='seo-master-2026'):
    """Main function to generate DOCX template from markdown"""
    parser = LabSubmissionParser(markdown_file)
    parsed_data = parser.parse()
    
    generator = DOCXGenerator(parsed_data)
    doc = generator.generate()
    
    # Generate output filename
    if output_dir:
        output_dir = Path(output_dir)
    else:
        # Default to course-specific assets directory
        output_dir = Path(f'data/courses/{course_name}/assets/templates')
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    day = parsed_data['metadata'].get('day', 0)
    lab = parsed_data['metadata'].get('lab_number', 0)
    filename = f"Day{day:02d}_Lab{lab:02d}_Submission_Template.docx"
    output_path = output_dir / filename
    
    doc.save(str(output_path))
    print(f"Generated: {output_path}")
    
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python parse_lab_submission.py <submission_markdown_file> [output_dir] [course_name]")
        sys.exit(1)
    
    markdown_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    course_name = sys.argv[3] if len(sys.argv) > 3 else 'seo-master-2026'
    
    try:
        generate_docx_template(markdown_file, output_dir, course_name)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

